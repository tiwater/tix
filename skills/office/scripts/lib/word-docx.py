#!/usr/bin/env python3
"""
word-docx.py — Cross-platform Word document backend using python-docx.

Works on Linux, Windows, and macOS without Microsoft Word installed.
Requires: pip install python-docx

Usage: python3 word-docx.py <function> [args...]
"""

import json
import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def require_docx():
    if not HAS_DOCX:
        print(json.dumps({
            "error": "python-docx not installed. Run: pip install python-docx"
        }))
        sys.exit(1)


def open_document(file_path=None):
    """Open a .docx file. If file_path is None, error (no active doc on Linux)."""
    require_docx()
    if not file_path:
        print(json.dumps({
            "error": "File path required on this platform (no active document support without Word)"
        }))
        sys.exit(1)
    if not os.path.exists(file_path):
        print(json.dumps({"error": f"File not found: {file_path}"}))
        sys.exit(1)
    return Document(file_path)


# ── Reading Functions ──

def readStructure(file_path=None):
    """Read document outline/structure."""
    doc = open_document(file_path)
    headings = []
    para_count = 0
    word_count = 0

    for i, para in enumerate(doc.paragraphs):
        para_count += 1
        word_count += len(para.text.split())
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading ", "").strip())
            except ValueError:
                level = 0
            headings.append({
                "index": i,
                "level": level,
                "text": para.text.strip()
            })

    print(json.dumps({
        "paragraphCount": para_count,
        "wordCount": word_count,
        "pageCount": "N/A (requires Word)",
        "headings": headings
    }, ensure_ascii=False))


def readSection(file_path=None, start=0, count=50):
    """Read a range of paragraphs."""
    doc = open_document(file_path)
    start = int(start)
    count = min(int(count), 200)
    paragraphs = []

    for i, para in enumerate(doc.paragraphs):
        if i < start:
            continue
        if i >= start + count:
            break
        paragraphs.append({
            "index": i,
            "style": para.style.name if para.style else "Normal",
            "text": para.text
        })

    print(json.dumps({
        "start": start,
        "count": len(paragraphs),
        "total": len(doc.paragraphs),
        "paragraphs": paragraphs
    }, ensure_ascii=False))


def readByHeading(file_path=None, heading_text=""):
    """Read all paragraphs under a specific heading."""
    doc = open_document(file_path)
    found = False
    found_level = 0
    paragraphs = []

    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ""

        if not found:
            if style_name.startswith("Heading") and heading_text.lower() in para.text.lower():
                found = True
                try:
                    found_level = int(style_name.replace("Heading ", "").strip())
                except ValueError:
                    found_level = 1
                paragraphs.append({
                    "index": i,
                    "style": style_name,
                    "text": para.text
                })
            continue

        # Stop at same-or-higher level heading
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading ", "").strip())
            except ValueError:
                level = 1
            if level <= found_level:
                break

        paragraphs.append({
            "index": i,
            "style": style_name,
            "text": para.text
        })

    print(json.dumps({
        "heading": heading_text,
        "found": found,
        "paragraphs": paragraphs
    }, ensure_ascii=False))


def readStyles(file_path=None):
    """List styles used in the document."""
    doc = open_document(file_path)
    style_counts = {}
    for para in doc.paragraphs:
        name = para.style.name if para.style else "Normal"
        style_counts[name] = style_counts.get(name, 0) + 1

    styles = [{"name": k, "count": v} for k, v in sorted(style_counts.items())]
    print(json.dumps({"styles": styles}, ensure_ascii=False))


def searchText(file_path=None, query=""):
    """Search for text in paragraphs."""
    doc = open_document(file_path)
    matches = []
    query_lower = query.lower()

    for i, para in enumerate(doc.paragraphs):
        if query_lower in para.text.lower():
            matches.append({
                "index": i,
                "style": para.style.name if para.style else "Normal",
                "text": para.text
            })
            if len(matches) >= 100:
                break

    print(json.dumps({
        "query": query,
        "matchCount": len(matches),
        "matches": matches
    }, ensure_ascii=False))


def getComments(file_path=None):
    """Extract comments from .docx (via XML parsing)."""
    require_docx()
    if not file_path:
        print(json.dumps({"error": "File path required"}))
        sys.exit(1)

    import zipfile
    from xml.etree import ElementTree as ET

    comments = []
    try:
        with zipfile.ZipFile(file_path) as zf:
            if "word/comments.xml" in zf.namelist():
                xml = zf.read("word/comments.xml")
                ns = {
                    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                }
                root = ET.fromstring(xml)
                for comment in root.findall(".//w:comment", ns):
                    text_parts = []
                    for p in comment.findall(".//w:t", ns):
                        if p.text:
                            text_parts.append(p.text)
                    comments.append({
                        "id": comment.get(f"{{{ns['w']}}}id", ""),
                        "author": comment.get(f"{{{ns['w']}}}author", ""),
                        "date": comment.get(f"{{{ns['w']}}}date", ""),
                        "text": " ".join(text_parts)
                    })
    except Exception as e:
        print(json.dumps({"error": str(e)}))
        sys.exit(1)

    print(json.dumps({
        "commentCount": len(comments),
        "comments": comments
    }, ensure_ascii=False))


# ── Creation Functions ──

def createDocument(title=None, save_path=None):
    """Create a new blank document."""
    require_docx()
    doc = Document()

    if title:
        doc.core_properties.title = title

    if save_path:
        doc.save(save_path)
        print(json.dumps({
            "created": True,
            "path": os.path.abspath(save_path),
            "title": title or ""
        }))
    else:
        print(json.dumps({
            "error": "save_path required on this platform (no active document support)"
        }))
        sys.exit(1)


def addHeading(file_path=None, text="", level=1):
    """Add a heading to the document."""
    doc = open_document(file_path)
    doc.add_heading(text, level=int(level))
    doc.save(file_path)
    print(json.dumps({"added": "heading", "text": text, "level": int(level)}))


def addParagraph(file_path=None, text="", bold=False, italic=False, style=None):
    """Add a paragraph to the document."""
    doc = open_document(file_path)
    para = doc.add_paragraph()
    if style:
        try:
            para.style = style
        except Exception:
            pass
    run = para.add_run(text)
    if bold or str(bold).lower() == "true":
        run.bold = True
    if italic or str(italic).lower() == "true":
        run.italic = True
    doc.save(file_path)
    print(json.dumps({"added": "paragraph", "text": text[:80]}))


def addTable(file_path=None, data_json=""):
    """Add a table from JSON data."""
    doc = open_document(file_path)
    data = json.loads(data_json)
    headers = data.get("headers", [])
    rows = data.get("rows", [])

    cols = len(headers) if headers else (len(rows[0]) if rows else 0)
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = "Table Grid"

    # Headers
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(header)
        for run in cell.paragraphs[0].runs:
            run.bold = True

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            table.rows[r + 1].cells[c].text = str(val)

    doc.save(file_path)
    print(json.dumps({
        "added": "table",
        "rows": len(rows),
        "cols": cols
    }))


def insertImage(file_path=None, image_path="", width=None, height=None):
    """Insert an image into the document."""
    doc = open_document(file_path)
    kwargs = {}
    if width:
        kwargs["width"] = Inches(float(width) / 72)  # points to inches
    if height:
        kwargs["height"] = Inches(float(height) / 72)
    doc.add_picture(image_path, **kwargs)
    doc.save(file_path)
    print(json.dumps({"added": "image", "path": image_path}))


def addPageBreak(file_path=None, break_type="page"):
    """Add a page or section break."""
    doc = open_document(file_path)
    if break_type == "section":
        doc.add_section()
    else:
        doc.add_page_break()
    doc.save(file_path)
    print(json.dumps({"added": f"{break_type}_break"}))


def addTableOfContents(file_path=None, levels=3):
    """Add a TOC placeholder (requires Word to update field codes)."""
    doc = open_document(file_path)
    para = doc.add_paragraph()
    para.text = "[Table of Contents — open in Word and press F9 to update]"
    para.style = "Normal"
    doc.save(file_path)
    print(json.dumps({
        "added": "toc_placeholder",
        "note": "Open in Word and press F9 to generate actual TOC"
    }))


def setHeaderFooter(file_path=None, position="header", text=""):
    """Set header or footer text."""
    doc = open_document(file_path)
    section = doc.sections[0]
    if position == "footer":
        section.footer.paragraphs[0].text = text
    else:
        section.header.paragraphs[0].text = text
    doc.save(file_path)
    print(json.dumps({"set": position, "text": text}))


def saveDocument(file_path=None, save_as=None, fmt=None):
    """Save/copy document. PDF export not supported without Word."""
    doc = open_document(file_path)
    target = save_as or file_path
    if fmt == "pdf":
        print(json.dumps({
            "error": "PDF export requires Microsoft Word. Save as .docx and convert separately."
        }))
        sys.exit(1)
    doc.save(target)
    print(json.dumps({"saved": True, "path": os.path.abspath(target)}))


def setDocumentProperties(file_path=None, **props):
    """Set document metadata properties."""
    doc = open_document(file_path)
    cp = doc.core_properties
    for key, val in props.items():
        if hasattr(cp, key):
            setattr(cp, key, val)
    doc.save(file_path)
    print(json.dumps({"set": list(props.keys())}))


# ── Review Functions (limited without Word) ──

def addComment(file_path=None, para_index=0, comment_text=""):
    """Add comment — limited on cross-platform (no native comment support in python-docx)."""
    print(json.dumps({
        "error": "Adding comments requires Microsoft Word. Use macOS or Windows backend.",
        "workaround": "Consider adding comments as highlighted text or annotations."
    }))
    sys.exit(1)


def addRevision(file_path=None, para_index=0, old_text="", new_text=""):
    """Add revision — limited on cross-platform."""
    print(json.dumps({
        "error": "Track changes requires Microsoft Word. Use macOS or Windows backend.",
        "workaround": "Consider making the edit directly with addParagraph."
    }))
    sys.exit(1)


# ── Dispatcher ──

FUNCTIONS = {
    "readStructure": readStructure,
    "readSection": readSection,
    "readByHeading": readByHeading,
    "readStyles": readStyles,
    "searchText": searchText,
    "getComments": getComments,
    "createDocument": createDocument,
    "addHeading": addHeading,
    "addParagraph": addParagraph,
    "addTable": addTable,
    "insertImage": insertImage,
    "addPageBreak": addPageBreak,
    "addTableOfContents": addTableOfContents,
    "setHeaderFooter": setHeaderFooter,
    "saveDocument": saveDocument,
    "setDocumentProperties": setDocumentProperties,
    "addComment": addComment,
    "addRevision": addRevision,
}

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"error": f"Usage: {sys.argv[0]} <function> [args...]"}))
        sys.exit(1)

    func_name = sys.argv[1]
    if func_name not in FUNCTIONS:
        print(json.dumps({"error": f"Unknown function: {func_name}", "available": list(FUNCTIONS.keys())}))
        sys.exit(1)

    args = sys.argv[2:]
    # Convert "null" string to None
    args = [None if a == "null" else a for a in args]
    FUNCTIONS[func_name](*args)
